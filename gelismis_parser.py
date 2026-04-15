"""
Gelişmiş Belge Okuyucu
========================
- PDF: Tablo algılama, yapı koruma, OCR desteği
- DOCX: Tablo ve başlık yapısını korur
- Görsel: OCR ile metin çıkarma
- Çıktı: Yapılandırılmış HTML veya temiz metin
"""
import re
import io
import logging
from pathlib import Path
from typing import Optional, List, Dict
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class BelgeBolum:
    """Bir belge bölümü (paragraf, tablo, başlık vb.)"""
    tip: str  # "baslik", "paragraf", "tablo", "madde", "liste", "gorsel_metin"
    icerik: str
    seviye: int = 0       # Başlık seviyesi (h1=1, h2=2...)
    sayfa_no: int = 0
    meta: dict = field(default_factory=dict)


@dataclass
class BelgeSonuc:
    """Belge okuma sonucu."""
    dosya_adi: str
    format: str
    bolumler: List[BelgeBolum]
    sayfa_sayisi: int = 0
    tablolar: List[dict] = field(default_factory=list)
    uyarilar: List[str] = field(default_factory=list)

    def to_html(self) -> str:
        """Yapılandırılmış HTML çıktısı üretir."""
        parcalar = [f'<div class="belge" data-dosya="{self.dosya_adi}">']
        for b in self.bolumler:
            if b.tip == "baslik":
                tag = f"h{min(b.seviye, 6)}" if b.seviye else "h2"
                parcalar.append(f'<{tag}>{_html_escape(b.icerik)}</{tag}>')
            elif b.tip == "tablo":
                parcalar.append(b.icerik)  # Zaten HTML tablo
            elif b.tip == "madde":
                parcalar.append(
                    f'<div class="madde" data-madde="{b.meta.get("no", "")}">'
                    f'<strong>Madde {b.meta.get("no", "")}</strong> – '
                    f'{_html_escape(b.icerik)}</div>'
                )
            elif b.tip == "liste":
                parcalar.append(f'<ul><li>{_html_escape(b.icerik)}</li></ul>')
            elif b.tip == "gorsel_metin":
                parcalar.append(
                    f'<div class="ocr-metin" data-sayfa="{b.sayfa_no}">'
                    f'<em>[OCR - Sayfa {b.sayfa_no}]</em><br>'
                    f'{_html_escape(b.icerik)}</div>'
                )
            else:
                if b.icerik.strip():
                    parcalar.append(f'<p>{_html_escape(b.icerik)}</p>')
        parcalar.append('</div>')
        return "\n".join(parcalar)

    def to_temiz_metin(self) -> str:
        """İndeksleme için temiz metin çıktısı."""
        parcalar = []
        for b in self.bolumler:
            if b.tip == "baslik":
                parcalar.append(f"\n{'#' * (b.seviye or 2)} {b.icerik}\n")
            elif b.tip == "tablo":
                parcalar.append(_tablo_html_to_metin(b.icerik))
            elif b.tip == "madde":
                parcalar.append(f"\nMadde {b.meta.get('no', '')}: {b.icerik}\n")
            else:
                if b.icerik.strip():
                    parcalar.append(b.icerik)
        return "\n".join(parcalar)


def _html_escape(metin: str) -> str:
    return (metin
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;"))


def _tablo_html_to_metin(html: str) -> str:
    """HTML tabloyu düz metin formatına çevirir."""
    import re
    # Satırları ayır
    satirlar = re.findall(r'<tr[^>]*>(.*?)</tr>', html, re.DOTALL)
    metin_satirlar = []
    for satir in satirlar:
        hucreler = re.findall(r'<t[dh][^>]*>(.*?)</t[dh]>', satir, re.DOTALL)
        temiz = [re.sub(r'<[^>]+>', '', h).strip() for h in hucreler]
        metin_satirlar.append(" | ".join(temiz))
    if metin_satirlar:
        # Başlık ayırıcı
        bolucu = "---|" * len(metin_satirlar[0].split("|"))
        metin_satirlar.insert(1, bolucu.rstrip("|"))
    return "\n".join(metin_satirlar)


# ══════════════════════════════════════════════════════════
#  PDF OKUYUCU (Gelişmiş)
# ══════════════════════════════════════════════════════════

def pdf_oku_gelismis(dosya_yolu: Path) -> BelgeSonuc:
    """
    PDF'i yapısal olarak okur:
    - Metin bloklarını paragraf/başlık olarak sınıflandırır
    - Tabloları HTML tabloya çevirir
    - Görsel sayfaları OCR ile okur
    """
    import fitz

    doc = fitz.open(str(dosya_yolu))
    bolumler: List[BelgeBolum] = []
    tablolar = []
    uyarilar = []
    metin_var = False

    for sayfa_idx, sayfa in enumerate(doc):
        sayfa_no = sayfa_idx + 1

        # ── Tabloları algıla ──
        try:
            sayfa_tablolari = sayfa.find_tables()
            if sayfa_tablolari and sayfa_tablolari.tables:
                for tablo in sayfa_tablolari.tables:
                    tablo_verisi = tablo.extract()
                    if tablo_verisi and len(tablo_verisi) > 1:
                        html_tablo = _tablo_to_html(tablo_verisi, sayfa_no)
                        bolumler.append(BelgeBolum(
                            tip="tablo",
                            icerik=html_tablo,
                            sayfa_no=sayfa_no,
                        ))
                        tablolar.append({
                            "sayfa": sayfa_no,
                            "satir": len(tablo_verisi),
                            "sutun": len(tablo_verisi[0]) if tablo_verisi else 0,
                        })
                        metin_var = True
        except Exception as e:
            logger.warning(f"Tablo algılama hatası (sayfa {sayfa_no}): {e}")

        # ── Metin bloklarını çıkar ──
        bloklar = sayfa.get_text("dict", sort=True).get("blocks", [])

        for blok in bloklar:
            if blok["type"] == 0:  # Metin bloğu
                blok_metin = ""
                max_font = 0
                is_bold = False

                for line in blok.get("lines", []):
                    satir_metin = ""
                    for span in line.get("spans", []):
                        satir_metin += span.get("text", "")
                        font_size = span.get("size", 12)
                        if font_size > max_font:
                            max_font = font_size
                        flags = span.get("flags", 0)
                        if flags & 2 ** 4:  # bold
                            is_bold = True
                    blok_metin += satir_metin.strip() + "\n"

                blok_metin = blok_metin.strip()
                if not blok_metin:
                    continue

                metin_var = True

                # Madde mi?
                madde_match = re.match(
                    r"(?:Madde|MADDE)\s+(\d+[\w/]*)\s*[-–:]?\s*(.*)",
                    blok_metin, re.DOTALL
                )
                if madde_match:
                    bolumler.append(BelgeBolum(
                        tip="madde",
                        icerik=madde_match.group(2).strip(),
                        sayfa_no=sayfa_no,
                        meta={"no": madde_match.group(1)},
                    ))
                    continue

                # Başlık mı? (büyük font veya kalın ve kısa)
                if (max_font > 14 or (is_bold and len(blok_metin) < 120)):
                    seviye = 1 if max_font > 18 else (2 if max_font > 14 else 3)
                    bolumler.append(BelgeBolum(
                        tip="baslik",
                        icerik=blok_metin,
                        seviye=seviye,
                        sayfa_no=sayfa_no,
                    ))
                    continue

                # Normal paragraf
                bolumler.append(BelgeBolum(
                    tip="paragraf",
                    icerik=blok_metin,
                    sayfa_no=sayfa_no,
                ))

            elif blok["type"] == 1:  # Görsel bloğu
                # Sayfa tamamıyla görsel mi kontrol edilecek
                pass

        # ── Metin yoksa OCR dene ──
        if not metin_var and sayfa_idx == 0:
            # İlk sayfada hiç metin yoksa muhtemelen taranmış belge
            ocr_metin = _sayfa_ocr(sayfa)
            if ocr_metin:
                bolumler.append(BelgeBolum(
                    tip="gorsel_metin",
                    icerik=ocr_metin,
                    sayfa_no=sayfa_no,
                ))
                uyarilar.append("Taranmış belge algılandı, OCR uygulandı.")
                metin_var = True

        # Sonraki sayfalarda da metin yoksa OCR
        sayfa_metni = sayfa.get_text().strip()
        if not sayfa_metni and sayfa_idx > 0:
            ocr_metin = _sayfa_ocr(sayfa)
            if ocr_metin:
                bolumler.append(BelgeBolum(
                    tip="gorsel_metin",
                    icerik=ocr_metin,
                    sayfa_no=sayfa_no,
                ))

    doc.close()

    return BelgeSonuc(
        dosya_adi=dosya_yolu.name,
        format="pdf",
        bolumler=bolumler,
        sayfa_sayisi=len(doc) if hasattr(doc, '__len__') else 0,
        tablolar=tablolar,
        uyarilar=uyarilar,
    )


def _tablo_to_html(veri: list, sayfa_no: int) -> str:
    """Tablo verisini HTML tabloya çevirir."""
    if not veri:
        return ""

    html = f'<table class="belge-tablo" data-sayfa="{sayfa_no}">\n'

    # İlk satır başlık
    html += "<thead><tr>"
    for hucre in veri[0]:
        html += f"<th>{_html_escape(str(hucre or ''))}</th>"
    html += "</tr></thead>\n"

    # Geri kalan satırlar
    html += "<tbody>\n"
    for satir in veri[1:]:
        html += "<tr>"
        for hucre in satir:
            html += f"<td>{_html_escape(str(hucre or ''))}</td>"
        html += "</tr>\n"
    html += "</tbody></table>"

    return html


def _sayfa_ocr(sayfa) -> str:
    """Sayfa görselini OCR ile okur."""
    try:
        import pytesseract
        from PIL import Image

        # Sayfayı yüksek çözünürlükte görsele çevir
        pix = sayfa.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))

        # Türkçe OCR
        metin = pytesseract.image_to_string(img, lang="tur+eng")
        return metin.strip() if metin.strip() else ""
    except ImportError:
        logger.warning("pytesseract veya PIL yüklü değil, OCR atlanıyor.")
        return ""
    except Exception as e:
        logger.warning(f"OCR hatası: {e}")
        return ""


# ══════════════════════════════════════════════════════════
#  DOCX OKUYUCU (Gelişmiş)
# ══════════════════════════════════════════════════════════

def docx_oku_gelismis(dosya_yolu: Path) -> BelgeSonuc:
    """DOCX'i yapısal olarak okur (başlıklar, tablolar, listeler)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(dosya_yolu))
    bolumler: List[BelgeBolum] = []
    tablolar = []

    # Paragrafları işle
    for para in doc.paragraphs:
        metin = para.text.strip()
        if not metin:
            continue

        stil = para.style.name.lower() if para.style else ""

        # Başlık mı?
        if "heading" in stil or "başlık" in stil:
            try:
                seviye = int(re.search(r'\d', stil).group())
            except (AttributeError, ValueError):
                seviye = 2
            bolumler.append(BelgeBolum(tip="baslik", icerik=metin, seviye=seviye))
            continue

        # Madde mi?
        madde_match = re.match(
            r"(?:Madde|MADDE)\s+(\d+[\w/]*)\s*[-–:]?\s*(.*)",
            metin, re.DOTALL
        )
        if madde_match:
            bolumler.append(BelgeBolum(
                tip="madde",
                icerik=madde_match.group(2).strip(),
                meta={"no": madde_match.group(1)},
            ))
            continue

        # Liste mi?
        if stil.startswith("list") or metin.startswith(("•", "-", "–", "►")):
            bolumler.append(BelgeBolum(tip="liste", icerik=metin.lstrip("•-–► ")))
            continue

        bolumler.append(BelgeBolum(tip="paragraf", icerik=metin))

    # Tabloları işle
    for tablo_idx, tablo in enumerate(doc.tables):
        veri = []
        for satir in tablo.rows:
            satir_veri = [hucre.text.strip() for hucre in satir.cells]
            veri.append(satir_veri)

        if veri and len(veri) > 1:
            html_tablo = _tablo_to_html(veri, 0)
            bolumler.append(BelgeBolum(tip="tablo", icerik=html_tablo))
            tablolar.append({
                "tablo_no": tablo_idx + 1,
                "satir": len(veri),
                "sutun": len(veri[0]) if veri else 0,
            })

    return BelgeSonuc(
        dosya_adi=dosya_yolu.name,
        format="docx",
        bolumler=bolumler,
        tablolar=tablolar,
    )


# ══════════════════════════════════════════════════════════
#  GÖRSEL OKUYUCU (OCR)
# ══════════════════════════════════════════════════════════

def gorsel_oku(dosya_yolu: Path) -> BelgeSonuc:
    """Görsel dosyayı OCR ile okur."""
    uyarilar = []
    bolumler = []

    try:
        import pytesseract
        from PIL import Image

        img = Image.open(str(dosya_yolu))

        # Ön işleme: gri tonlama, kontrast artırma
        if img.mode != "L":
            img = img.convert("L")

        metin = pytesseract.image_to_string(img, lang="tur+eng")

        if metin.strip():
            # Madde yapısını bulmaya çalış
            satirlar = metin.strip().split("\n")
            mevcut_paragraf = []

            for satir in satirlar:
                satir = satir.strip()
                if not satir:
                    if mevcut_paragraf:
                        bolumler.append(BelgeBolum(
                            tip="paragraf",
                            icerik=" ".join(mevcut_paragraf),
                        ))
                        mevcut_paragraf = []
                    continue

                madde_match = re.match(
                    r"(?:Madde|MADDE)\s+(\d+[\w/]*)\s*[-–:]?\s*(.*)",
                    satir
                )
                if madde_match:
                    if mevcut_paragraf:
                        bolumler.append(BelgeBolum(
                            tip="paragraf",
                            icerik=" ".join(mevcut_paragraf),
                        ))
                        mevcut_paragraf = []
                    bolumler.append(BelgeBolum(
                        tip="madde",
                        icerik=madde_match.group(2).strip(),
                        meta={"no": madde_match.group(1)},
                    ))
                else:
                    mevcut_paragraf.append(satir)

            if mevcut_paragraf:
                bolumler.append(BelgeBolum(
                    tip="paragraf",
                    icerik=" ".join(mevcut_paragraf),
                ))
        else:
            uyarilar.append("Görselden metin çıkarılamadı.")

    except ImportError:
        uyarilar.append("OCR için pytesseract ve Pillow gereklidir.")
    except Exception as e:
        uyarilar.append(f"Görsel okuma hatası: {str(e)}")

    return BelgeSonuc(
        dosya_adi=dosya_yolu.name,
        format="gorsel",
        bolumler=bolumler,
        uyarilar=uyarilar,
    )


# ══════════════════════════════════════════════════════════
#  TXT OKUYUCU (Akıllı Yapı Algılama)
# ══════════════════════════════════════════════════════════

def txt_oku_gelismis(dosya_yolu: Path) -> BelgeSonuc:
    """TXT dosyasını akıllıca yapılandırır."""
    metin = ""
    for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1254"):
        try:
            metin = dosya_yolu.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, Exception):
            continue

    if not metin:
        return BelgeSonuc(
            dosya_adi=dosya_yolu.name, format="txt",
            bolumler=[], uyarilar=["Dosya okunamadı."],
        )

    bolumler = []
    satirlar = metin.split("\n")
    mevcut_paragraf = []

    for satir in satirlar:
        satir_strip = satir.strip()

        if not satir_strip:
            if mevcut_paragraf:
                bolumler.append(BelgeBolum(
                    tip="paragraf",
                    icerik=" ".join(mevcut_paragraf),
                ))
                mevcut_paragraf = []
            continue

        # Tamamen büyük harf → başlık
        if (satir_strip.isupper() and len(satir_strip) > 3
                and len(satir_strip) < 150):
            if mevcut_paragraf:
                bolumler.append(BelgeBolum(
                    tip="paragraf",
                    icerik=" ".join(mevcut_paragraf),
                ))
                mevcut_paragraf = []
            bolumler.append(BelgeBolum(
                tip="baslik", icerik=satir_strip, seviye=2,
            ))
            continue

        # Madde
        madde_match = re.match(
            r"(?:Madde|MADDE)\s+(\d+[\w/]*)\s*[-–:]?\s*(.*)",
            satir_strip, re.DOTALL,
        )
        if madde_match:
            if mevcut_paragraf:
                bolumler.append(BelgeBolum(
                    tip="paragraf",
                    icerik=" ".join(mevcut_paragraf),
                ))
                mevcut_paragraf = []
            bolumler.append(BelgeBolum(
                tip="madde",
                icerik=madde_match.group(2).strip(),
                meta={"no": madde_match.group(1)},
            ))
            continue

        mevcut_paragraf.append(satir_strip)

    if mevcut_paragraf:
        bolumler.append(BelgeBolum(
            tip="paragraf",
            icerik=" ".join(mevcut_paragraf),
        ))

    return BelgeSonuc(
        dosya_adi=dosya_yolu.name, format="txt", bolumler=bolumler,
    )


# ══════════════════════════════════════════════════════════
#  ANA GİRİŞ NOKTASI
# ══════════════════════════════════════════════════════════

GORSEL_UZANTILAR = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}

def belge_isle(dosya_yolu: str | Path) -> BelgeSonuc:
    """
    Dosya formatına göre en uygun okuyucuyu seçer.
    Yapılandırılmış BelgeSonuc döndürür.
    """
    yol = Path(dosya_yolu)
    if not yol.exists():
        return BelgeSonuc(
            dosya_adi=yol.name, format="?", bolumler=[],
            uyarilar=["Dosya bulunamadı."],
        )

    uzanti = yol.suffix.lower()

    if uzanti == ".pdf":
        return pdf_oku_gelismis(yol)
    elif uzanti in (".docx", ".doc"):
        return docx_oku_gelismis(yol)
    elif uzanti in GORSEL_UZANTILAR:
        return gorsel_oku(yol)
    elif uzanti in (".txt", ".md", ".csv"):
        return txt_oku_gelismis(yol)
    else:
        return BelgeSonuc(
            dosya_adi=yol.name, format=uzanti,
            bolumler=[], uyarilar=[f"Desteklenmeyen format: {uzanti}"],
        )
